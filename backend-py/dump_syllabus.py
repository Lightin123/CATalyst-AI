import docx

doc = docx.Document('Syllabus.docx')
with open('syllabus_dump.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text.strip() + '\n')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    f.write(text + '\n')
